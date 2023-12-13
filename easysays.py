import requests
from bs4 import BeautifulSoup
from docx import Document
import os
import logging
from concurrent.futures import ThreadPoolExecutor

# Set up basic logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def fetch_page(url):
    try:
        response = requests.get(url)
        response.raise_for_status()
        return response.text
    except requests.RequestException as e:
        logging.error(f"Error fetching {url}: {e}")
        return None

def parse_main_page(content):
    soup = BeautifulSoup(content, 'html.parser')
    links = soup.find_all('a')
    article_links = [link['href'] for link in links if 'html' in link['href']]
    return ["http://paulgraham.com/" + link for link in article_links if not link.startswith("http")]

def fetch_and_parse_article(url):
    article_content = fetch_page(url)
    if article_content:
        title, content = parse_article(article_content)
        if title and content:
            return title, content
        else:
            logging.error(f"Failed to parse article: {url}")
    else:
        logging.error(f"Failed to fetch article: {url}")
    return None, None

def parse_article(content):
    try:
        soup = BeautifulSoup(content, 'html.parser')
        title = soup.find('title').text
        body = soup.find('body')
        if not body:
            raise ValueError("Body tag not found")
        article_content = body.text.strip()
        return title, article_content
    except Exception as e:
        logging.error(f"Error parsing article: {e}")
        return None, None

def add_article_to_docx(doc, title, content):
    doc.add_heading(title, level=1)
    doc.add_paragraph(content)
    doc.add_page_break()

def main():
    os.makedirs('articles', exist_ok=True)
    main_page_url = "http://paulgraham.com/articles.html"
    main_page_content = fetch_page(main_page_url)

    if main_page_content:
        article_links = parse_main_page(main_page_content)

        doc = Document()
        doc.add_heading('Collected Articles from Paul Graham', 0)

        with ThreadPoolExecutor(max_workers=10) as executor:
            results = executor.map(fetch_and_parse_article, article_links)

            for result in results:
                if result:
                    title, content = result
                    add_article_to_docx(doc, title, content)
                    logging.info(f"Added '{title}' to the document.")

        doc.save(os.path.join('articles', 'Collected_Articles.docx'))
        logging.info("All articles saved in a book format to a Word document.")

if __name__ == "__main__":
    main()